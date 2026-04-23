import os
import io
import uuid
import pandas as pd
import urllib.parse

# Global dictionary to store dataframes safely in memory
# Structure: { "session_id": { "file_name_sheet_name": df } }
DATASETS = {}
SESSION_FILES = {}

def export_dataframe(df: pd.DataFrame, file_format: str, base_filename: str):
    os.makedirs("exports", exist_ok=True)
    file_id = str(uuid.uuid4())[:8]
    filepath = f"exports/{base_filename}_{file_id}.{file_format}"
    
    if file_format == "xlsx":
        df.to_excel(filepath, index=False)
    elif file_format == "csv":
        df.to_csv(filepath, index=False)
    elif file_format == "pdf":
        try:
            from fpdf import FPDF
            orientation = 'P' if len(df.columns) <= 4 else 'L'
            pdf = FPDF(orientation=orientation)
            pdf.add_page()
            pdf.set_font("Arial", size=8)
            
            # Calculate proportional column widths
            col_widths = []
            for col in df.columns:
                max_len = max(df[col].astype(str).map(len).max() if not df.empty else 0, len(str(col)))
                col_widths.append(min(max_len * 1.5 + 4, 60))
                
            # Scale to page width
            usable_width = pdf.w - 2 * pdf.l_margin
            total_width = sum(col_widths)
            if total_width > usable_width:
                scale = usable_width / total_width
                col_widths = [w * scale for w in col_widths]
                
            row_height = pdf.font_size * 1.5
            
            # Header
            for i, col in enumerate(df.columns):
                trunc_len = max(int(col_widths[i] / 1.5), 1)
                pdf.cell(col_widths[i], row_height, txt=str(col)[:trunc_len], border=1)
            pdf.ln(row_height)
            
            # Data
            for _, row in df.head(200).iterrows(): # Limit for PDF readability
                for i, val in enumerate(row):
                    trunc_len = max(int(col_widths[i] / 1.5), 1)
                    pdf.cell(col_widths[i], row_height, txt=str(val)[:trunc_len], border=1)
                pdf.ln(row_height)
                
            pdf.output(filepath)
        except Exception as e:
            return f"Error creating PDF: {str(e)}"
    elif file_format == "docx":
        try:
            from docx import Document
            from docx.enum.section import WD_ORIENT
            
            doc = Document()
            
            # Dynamic orientation based on column count
            if len(df.columns) > 4:
                section = doc.sections[-1]
                new_width, new_height = section.page_height, section.page_width
                section.orientation = WD_ORIENT.LANDSCAPE
                section.page_width = new_width
                section.page_height = new_height
            
            doc.add_heading(f"{base_filename} Data Export", 0)
            
            export_df = df.head(500) # Increased docx capacity
            t = doc.add_table(export_df.shape[0]+1, export_df.shape[1])
            t.style = 'Table Grid'
            t.autofit = True
            
            # Header
            for j in range(export_df.shape[-1]):
                t.cell(0,j).text = str(export_df.columns[j])

            # Data
            for i in range(export_df.shape[0]):
                for j in range(export_df.shape[-1]):
                    t.cell(i+1,j).text = str(export_df.values[i,j])
            
            doc.save(filepath)
        except Exception as e:
            return f"Error creating Word doc: {str(e)}"
    else:
        return f"Error: Unsupported format {file_format}"
        
    # Return a preview of the data and a markdown link to download the file
    preview = df.to_string(index=False)
    encoded_filepath = urllib.parse.quote(filepath, safe='/')
    return f"Extracted Data Preview:\n{preview}\n\n[Download {file_format.upper()} File](/{encoded_filepath})"

def clean_headers(df):
    # Drop completely empty rows and columns
    df.dropna(how='all', inplace=True)
    df.dropna(axis=1, how='all', inplace=True)
    
    # If the dataframe has Unnamed columns, the actual header might be the first row
    unnamed_cols = [c for c in df.columns if "Unnamed" in str(c)]
    if len(unnamed_cols) > 0 and not df.empty:
        # If more than half the columns are Unnamed, promote the first row
        if len(unnamed_cols) >= len(df.columns) / 2:
            new_header = df.iloc[0]
            df = df[1:]
            df.columns = new_header
            df.reset_index(drop=True, inplace=True)
            
            # Drop any columns that became NaN after promotion
            df = df.loc[:, df.columns.notna()]
    
    # Clean up column names (convert to string, strip whitespace)
    df.columns = [str(c).strip() for c in df.columns]
    return df

def process_upload(file_contents: bytes, filename: str, session_id: str):
    if session_id not in DATASETS:
        DATASETS[session_id] = {}
    if session_id not in SESSION_FILES:
        SESSION_FILES[session_id] = set()
        
    SESSION_FILES[session_id].add(filename)
    session_datasets = DATASETS[session_id]
    
    # Save file to disk persistently
    os.makedirs("uploads", exist_ok=True)
    file_path = os.path.join("uploads", filename)
    with open(file_path, "wb") as f:
        f.write(file_contents)
    
    filename_lower = filename.lower()
    if filename_lower.endswith(".csv"):
        try:
            df = pd.read_csv(io.BytesIO(file_contents), encoding='utf-8')
        except UnicodeDecodeError:
            df = pd.read_csv(io.BytesIO(file_contents), encoding='cp1252')
            
        df = clean_headers(df)
            
        dataset_name = filename.rsplit(".", 1)[0]
        session_datasets[dataset_name] = df
        return [dataset_name]
        
    elif filename_lower.endswith(".xlsx") or filename_lower.endswith(".xls"):
        xls = pd.read_excel(io.BytesIO(file_contents), sheet_name=None)
        dataset_names = []
        for sheet_name, df in xls.items():
            df = clean_headers(df)
            
            dataset_name = f"{filename.rsplit('.', 1)[0]}_{sheet_name}"
            session_datasets[dataset_name] = df
            dataset_names.append(dataset_name)
        return dataset_names
    else:
        raise ValueError("Unsupported file format. Only CSV and Excel files are allowed.")

def get_global_files():
    os.makedirs("uploads", exist_ok=True)
    files = os.listdir("uploads")
    # Return valid only
    return [f for f in files if f.endswith(".csv") or f.endswith(".xlsx") or f.endswith(".xls")]

def use_global_file(filename: str, session_id: str):
    file_path = os.path.join("uploads", filename)
    if not os.path.exists(file_path):
        raise ValueError("File not found")
    
    with open(file_path, "rb") as f:
        contents = f.read()
        
    return process_upload(contents, filename, session_id)

def get_schema_summary(session_id: str):
    if session_id not in DATASETS or not DATASETS[session_id]:
        return "No datasets available."
        
    summary = []
    for name, df in DATASETS[session_id].items():
        columns = df.dtypes.to_dict()
        col_str = ", ".join([f"'{col}': {dtype}" for col, dtype in columns.items()])
        sample = df.head(5).to_dict(orient="records")
        total_rows = len(df)
        summary.append(
            f"Dataset Name: `{name}`\n"
            f"Total Rows: {total_rows} (the dfs dictionary contains ALL {total_rows} rows)\n"
            f"Columns: {col_str}\n"
            f"Sample Data (first 5 rows only — for reference): {sample}\n"
        )
        
    return "\n".join(summary)

def execute_pandas_code(code: str, session_id: str):
    # Security: Stop execution of arbitrary unsafe commands
    disallowed = ["import os", "import sys", "subprocess", "eval(", "open("]
    if any(kw in code for kw in disallowed):
        return "Error: Unsafe code detected. Execution blocked."
        
    if session_id not in DATASETS:
        return "Error: No data available for this session."
        
    global_env = {
        "pd": pd,
        "dfs": DATASETS[session_id], # dict of dataframes
        "export_dataframe": export_dataframe
    }
    
    try:
        # We enforce code to provide 'result'
        exec(code, global_env)
        if 'result' in global_env:
            res = global_env['result']
            if isinstance(res, pd.DataFrame):
                return res.to_json(orient="records")
            elif isinstance(res, pd.Series):
                return res.to_json()
            else:
                return str(res)
        else:
            return "Error: The AI code did not define a 'result' variable."
    except Exception as e:
        return f"Execution Error: {str(e)}"
